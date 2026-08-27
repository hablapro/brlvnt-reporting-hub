#!/usr/bin/env python3
"""Apply the brlvnt house style to a markdown document and emit a formatted .docx.

Why this exists: pandoc's default conversion is Cambria headings and hairline
tables. It is a format change, not formatting, and shipping it as a client
deliverable was a standing gap through the August 2026 GGMI cycle. This is the
document-side equivalent of lib/housestyle.py, which does the same job for decks.

Tokens come from docs/DESIGN-SYSTEM.md. Do not declare a colour or a font here
that is not in that file.

    python3 scripts/format_doc.py <input.md> <output.docx>

Run it on every client-facing document before upload. There is no separate
restyle pass and no "good enough" default.
"""
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ---- tokens (docs/DESIGN-SYSTEM.md) ----
NAVY = RGBColor(0x1E, 0x27, 0x61)
DEEP = RGBColor(0x0F, 0x15, 0x35)
CORAL = RGBColor(0xF9, 0x61, 0x67)
INK = RGBColor(0x2B, 0x31, 0x47)
MUTED = RGBColor(0x5A, 0x60, 0x72)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEAD_FILL = "1E2761"
ALT_FILL = "F9FAFC"
TOTAL_FILL = "CADCFC"
BORDER = "D0D4DC"

SERIF, SANS = "Georgia", "Calibri"
# Heading level -> (pt, colour, space_before_pt)
H = {1: (17, NAVY, 20), 2: (13, NAVY, 16), 3: (11.5, NAVY, 12), 4: (10.5, MUTED, 10)}


def _shade(cell, hexfill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def _borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), BORDER)
        borders.append(e)
    tblPr.append(borders)


def _numericish(text):
    t = text.strip().replace("$", "").replace(",", "").replace("%", "")
    t = t.replace("~", "").replace("+", "").replace("-", "").replace("*", "")
    return bool(t) and t.replace(".", "", 1).isdigit()


def style_document(path):
    doc = Document(path)

    # Base font for everything, so stray runs do not fall back to Cambria.
    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), SANS)

    for p in doc.paragraphs:
        name = (p.style.name or "").lower()
        if name.startswith("title"):
            for r in p.runs:
                r.font.name, r.font.size, r.font.bold = SERIF, Pt(22), True
                r.font.color.rgb = DEEP
            p.paragraph_format.space_after = Pt(4)
        elif name.startswith("heading"):
            try:
                lvl = int(name.split()[-1])
            except ValueError:
                lvl = 3
            size, colour, before = H.get(lvl, H[4])
            for r in p.runs:
                r.font.name, r.font.size, r.font.bold = SERIF, Pt(size), True
                r.font.color.rgb = colour
            p.paragraph_format.space_before = Pt(before)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
        else:
            for r in p.runs:
                if not r.font.name:
                    r.font.name = SANS
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _borders(table)
        for i, row in enumerate(table.rows):
            # Never let a row break across a page, and repeat the header on
            # any table that does span one.
            trPr = row._tr.get_or_add_trPr()
            trPr.append(OxmlElement("w:cantSplit"))
            if i == 0:
                trPr.append(OxmlElement("w:tblHeader"))
            for cell in row.cells:
                cell.vertical_alignment = 1
                txt = cell.text
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    if i > 0 and _numericish(txt):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for r in p.runs:
                        r.font.name = SANS
                        r.font.size = Pt(9.5)
                        if i == 0:
                            r.font.bold = True
                            r.font.color.rgb = WHITE
                        else:
                            r.font.color.rgb = INK
            if i == 0:
                for cell in row.cells:
                    _shade(cell, HEAD_FILL)
            else:
                # A summary row is bold all the way across. A data row with a
                # bolded label or a bolded figure is not, and must not take the
                # total fill (this caught Bing and Azerion in the August draft).
                def cell_bold(c):
                    runs = [r for pp in c.paragraphs for r in pp.runs if r.text.strip()]
                    return bool(runs) and all(r.bold for r in runs)
                cells = [c for c in row.cells if c.text.strip()]
                is_total = bool(cells) and all(cell_bold(c) for c in cells)
                fill = TOTAL_FILL if is_total else (ALT_FILL if i % 2 == 0 else None)
                if fill:
                    for cell in row.cells:
                        _shade(cell, fill)

    doc.save(path)
    return len(doc.tables), len(doc.paragraphs)


def main():
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    src, dst = Path(sys.argv[1]), Path(sys.argv[2])
    subprocess.run(
        ["pandoc", str(src), "-o", str(dst), "--from=gfm"], check=True
    )
    tables, paras = style_document(dst)
    print(f"formatted {dst.name}: {tables} tables, {paras} paragraphs, house style applied")


if __name__ == "__main__":
    main()
